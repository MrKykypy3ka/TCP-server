from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from decimal import Decimal
import pandas as pd


def save_tables_and_images_to_word(table1, table2, table3, table4, image1, image2, output_file):
    # Создание документа Word
    doc = Document()

    # Установка полей документа
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Стиль текста
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет

    # Стиль заголовков
    heading_style = doc.styles.add_style('HeadingStyle', 1)
    heading_font = heading_style.font
    heading_font.name = 'Times New Roman'
    heading_font.size = Pt(14)
    heading_font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    heading_font.bold = True  # Полужирный шрифт


    doc.add_heading('Текущая ситуация на дороге дорог без учета временных промежутков и дней недели',
                    level=1).style = heading_style
    # Создание таблицы
    num_rows = len(table1.index) + 1
    num_cols = len(table1.columns) + 1
    table_docx = doc.add_table(rows=num_rows, cols=num_cols)

    # Заполнение заголовков столбцов
    header_cells = table_docx.rows[0].cells
    header_cells[0].text = "type_of_road"
    for j, column_name in enumerate(table1.columns):
        header_cells[j + 1].text = str(column_name)

    # Заполнение ячеек таблицы
    for i, row_name in enumerate(table1.index):
        row_cells = table_docx.rows[i + 1].cells
        row_cells[0].text = str(row_name)
        for j, column_name in enumerate(table1.columns):
            value = table1.loc[row_name, column_name]
            if not pd.isnull(value):
                row_cells[j + 1].text = str(value)
            else:
                row_cells[j + 1].text = '-'

    # Заголовок 2 и таблица 2
    doc.add_paragraph()
    doc.add_heading('Ситуация на дороге после перекрытия дорог без учета временных промежутков и дней недели', level=1).style = heading_style

    # Создание таблицы
    num_rows = len(table2.index) + 1
    num_cols = len(table2.columns) + 1
    table_docx = doc.add_table(rows=num_rows, cols=num_cols)

    # Заполнение заголовков столбцов
    header_cells = table_docx.rows[0].cells
    header_cells[0].text = "type_of_road"
    for j, column_name in enumerate(table2.columns):
        header_cells[j + 1].text = str(column_name)

    # Заполнение ячеек таблицы
    for i, row_name in enumerate(table2.index):
        row_cells = table_docx.rows[i + 1].cells
        row_cells[0].text = str(row_name)
        for j, column_name in enumerate(table2.columns):
            value = table2.loc[row_name, column_name]
            if not pd.isnull(value):
                row_cells[j + 1].text = str(value)
            else:
                row_cells[j + 1].text = '-'


    # Заголовок 3 и таблица 3
    doc.add_paragraph()
    doc.add_heading('Текущая ситуация на дороге c учета временных промежутков и дней недели', level=1).style = heading_style
    table = doc.add_table(rows=1, cols=len(table3.columns) + 1)

    # Заполняем заголовки столбцов
    table.cell(0, 0).text = 'date'
    for i, column in enumerate(table3.columns):
        table.cell(0, i + 1).text = str(column)

    # Заполняем данные таблицы
    for i, (index, row) in enumerate(table3.iterrows()):
        table.add_row()
        table.cell(i + 1, 0).text = str(index)
        for j, value in enumerate(row):
            table.cell(i + 1, j + 1).text = str(value)

    # Заголовок 4 и таблица 4
    doc.add_paragraph()
    doc.add_heading('Ситуация на дороге после перекрытия дорог c учета временных промежутков и дней недели',
                    level=1).style = heading_style
    table = doc.add_table(rows=1, cols=len(table4.columns) + 1)

    # Заполняем заголовки столбцов
    table.cell(0, 0).text = 'date'
    for i, column in enumerate(table4.columns):
        table.cell(0, i + 1).text = str(column)

    # Заполняем данные таблицы
    for i, (index, row) in enumerate(table4.iterrows()):
        table.add_row()
        table.cell(i + 1, 0).text = str(index)
        for j, value in enumerate(row):
            table.cell(i + 1, j + 1).text = str(value)

    # Вставка изображения 1
    doc.add_paragraph()
    doc.add_heading('Ситуация на дороге после перекрытия без учёта временных промежутков и дней недели', level=1).style = heading_style
    doc.add_picture(image1, width=Cm(15), height=Cm(8))

    # Вставка изображения 2
    doc.add_paragraph()
    doc.add_heading('Ситуация на дороге после перекрытия дорог с учетом дней недели', level=1).style = heading_style
    doc.add_picture(image2, width=Cm(19.5), height=Cm(12.5))

    # Сохранение файла Word
    doc.save(output_file)
